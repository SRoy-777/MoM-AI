import os
import io
import json
import logging
from typing import Dict, Any, List, Optional

# Auto-load .env file if present
if os.path.exists(".env"):
    with open(".env", "r") as f:
        for line in f:
            if "=" in line and not line.startswith("#"):
                k, v = line.strip().split("=", 1)
                os.environ.setdefault(k.strip(), v.strip())

from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from services.gemini_mom import GeminiMoMService
from services.transcription import AudioTranscriptionService
from services.bot_service import TeamsBotService

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("mom_ai")

app = FastAPI(
    title="MoM AI - Free AI Meeting Assistant & MoM Generator",
    description="Captures unrecorded 4-hour MS Teams meetings, integrates human assistant notes, and generates executive Minutes of Meeting.",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Services
gemini_service = GeminiMoMService()
transcription_service = AudioTranscriptionService()
bot_service = TeamsBotService()

class TaskItem(BaseModel):
    task: str
    assignee: Optional[str] = "Unassigned"
    department: Optional[str] = "General"
    timeframe: Optional[str] = "ASAP"
    priority: Optional[str] = "Medium"
    status: Optional[str] = "Pending"

class MoMRequest(BaseModel):
    transcript: str
    human_notes: Optional[str] = ""
    manual_tasks: Optional[List[TaskItem]] = []
    meeting_title: Optional[str] = "Executive MoM"
    meeting_date: Optional[str] = ""
    attendees: Optional[str] = ""
    api_key: Optional[str] = ""

@app.post("/api/generate-mom")
async def generate_mom_endpoint(req: MoMRequest):
    try:
        api_key = req.api_key or os.environ.get("GEMINI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=400, detail="Gemini API Key is required. Please provide it in the input settings or set GEMINI_API_KEY environment variable.")

        gemini_service.update_api_key(api_key)

        tasks_dict = [t.model_dump() for t in req.manual_tasks] if req.manual_tasks else []

        mom_data = gemini_service.generate_mom(
            transcript=req.transcript,
            human_notes=req.human_notes or "",
            manual_tasks=tasks_dict,
            meeting_title=req.meeting_title or "Meeting Minutes",
            meeting_date=req.meeting_date or "",
            attendees=req.attendees or ""
        )

        return {"success": True, "mom": mom_data}
    except Exception as e:
        logger.error(f"Error generating MoM: {e}")
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})

@app.post("/api/transcribe")
async def transcribe_endpoint(file: UploadFile = File(...)):
    try:
        content = await file.read()
        res = transcription_service.transcribe_audio_bytes(content, filename=file.filename)
        return res
    except Exception as e:
        logger.error(f"Upload transcription error: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/teams-bot/join")
async def join_bot_endpoint(meeting_url: str = Form(...), bot_name: str = Form("MoM AI Note Taker")):
    res = bot_service.join_teams_meeting(meeting_url, bot_name)
    return res

@app.get("/api/teams-bot/status")
async def bot_status_endpoint(session_id: str = ""):
    return bot_service.get_status(session_id)

@app.post("/api/teams-bot/leave")
async def leave_bot_endpoint(session_id: str = Form(...)):
    res = bot_service.leave_teams_meeting(session_id)
    return res

@app.post("/api/export-docx")
async def export_docx(req: Dict[str, Any]):
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor

        doc = docx.Document()
        title_p = doc.add_paragraph()
        run = title_p.add_run(req.get("title", "Minutes of Meeting"))
        run.bold = True
        run.font.size = Pt(22)

        doc.add_paragraph(f"Date: {req.get('date', 'N/A')}")
        doc.add_paragraph(f"Attendees: {req.get('attendees', 'N/A')}")

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(req.get("executive_summary", ""))

        doc.add_heading("Discussion Highlights", level=1)
        for topic in req.get("discussion_highlights", []):
            doc.add_heading(topic.get("topic", "Topic"), level=2)
            for pt in topic.get("points", []):
                doc.add_paragraph(f"• {pt}")

        doc.add_heading("Key Decisions", level=1)
        for dec in req.get("decisions", []):
            doc.add_paragraph(f"✔ {dec}")

        doc.add_heading("Action Items Matrix", level=1)
        actions = req.get("action_items", [])
        if actions:
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Task'
            hdr_cells[1].text = 'Assignee'
            hdr_cells[2].text = 'Department'
            hdr_cells[3].text = 'Timeframe'
            hdr_cells[4].text = 'Priority'

            for item in actions:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get("task", "")
                row_cells[1].text = item.get("assignee", "")
                row_cells[2].text = item.get("department", "")
                row_cells[3].text = item.get("timeframe", "")
                row_cells[4].text = item.get("priority", "")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=MoM_Report.docx"}
        )
    except Exception as e:
        logger.error(f"Docx export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Mount static files
static_dir = os.path.join(os.path.dirname(__file__), "static")
if not os.path.exists(static_dir):
    os.makedirs(static_dir)

app.mount("/static", StaticFiles(directory=static_dir), name="static")

@app.get("/", response_class=HTMLResponse)
async def read_index():
    index_path = os.path.join(static_dir, "index.html")
    if os.path.exists(index_path):
        with open(index_path, "r", encoding="utf-8") as f:
            return f.read()
    return "<h1>MoM AI Backend Running</h1>"

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=7860, reload=True)
